import json
import os
import docx
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from fpdf import FPDF
from crypto import *
import fitz  # PyMuPDF


def json_to_word(json_data, word_file):
    # 创建 Word 文档
    doc = Document()

    # 设置文档样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'  # 修改为常见中文字体
    font.size = Pt(12)

    # 添加标题
    doc.add_heading('医疗记录', 0)

    # 遍历 JSON 数据并添加到 Word 文档
    for record in json_data:
        # 添加记录标题
        doc.add_heading(f'记录ID: {record["medical_record_id"]}', level=1)

        # 添加患者信息
        doc.add_heading('基本信息', level=2)
        doc.add_paragraph(f'患者ID: {record["patient_id"]}')
        doc.add_paragraph(f'医生ID: {record["doctor_id"]}')
        doc.add_paragraph(f'就诊日期: {record["visit_date"]}')
        doc.add_paragraph(f'科室: {record["department"]}')

        # 添加医疗记录详细信息
        doc.add_heading('医疗记录详细信息', level=2)
        doc.add_paragraph(f'患者主诉: {record["patient_complaint"]}')
        doc.add_paragraph(f'病史: {record["medical_history"]}')
        doc.add_paragraph(f'体格检查: {record["physical_examination"]}')
        doc.add_paragraph(f'辅助检查: {record["auxiliary_examination"]}')
        doc.add_paragraph(f'诊断: {record["diagnosis"]}')
        doc.add_paragraph(f'治疗建议: {record["treatment_advice"]}')

        # 添加签名信息
        doc.add_heading('签名信息', level=2)
        doc.add_paragraph(f'医生签名: {record["doctor_signature"]}')
        doc.add_paragraph(f'创建时间: {record["created_at"]}')
        doc.add_paragraph(f'医生公钥: {record["doctor_public_key"]}')
        doc.add_paragraph(f'验证字符串: {record["to_cal"]}')
        doc.add_paragraph(f'验证签名: {record["to_verify_signature"]}')

        # 添加分页符
        doc.add_page_break()

        # 设置汉字字体
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 保存 Word 文档
    doc.save(word_file)
    print(f"Word 文档已保存为: {word_file}")

def json_to_pdf(json_data, pdf_file):
    # 创建 PDF 文档
    pdf = FPDF()
    pdf.add_page()
    current_dir = os.path.dirname(__file__)  # 获取当前脚本所在目录
    font_path = os.path.join(current_dir, "simsun.ttf")  # 指定字体文件路径

    pdf.add_font("SimSun", style="", fname=font_path, uni=True)  # 添加宋体常规字体
    pdf.set_font("SimSun", size=10)

    # 添加标题
    pdf.set_font("SimSun", '', 14)
    pdf.cell(200, 10, txt="医疗记录", ln=True, align='C')

    # 遍历 JSON 数据并添加到 PDF 文档
    for record in json_data:
        # 添加记录标题
        pdf.set_font("SimSun", '', 10)
        pdf.cell(200, 10, txt=f'病历ID: {record["medical_record_id"]}', ln=True)

        # 添加患者信息
        pdf.set_font("SimSun", '', 12)
        pdf.cell(200, 10, txt='基本信息', ln=True)
        pdf.set_font("SimSun", size=10)
        pdf.cell(200, 10, txt=f'患者ID: {record["patient_id"]}', ln=True)
        pdf.cell(200, 10, txt=f'医生ID: {record["doctor_id"]}', ln=True)
        pdf.cell(200, 10, txt=f'就诊日期: {record["visit_date"]}', ln=True)
        pdf.cell(200, 10, txt=f'科室: {record["department"]}', ln=True)

        # 添加医疗记录详细信息
        pdf.set_font("SimSun", '', 12)
        pdf.cell(200, 10, txt='医疗记录详细信息', ln=True)
        pdf.set_font("SimSun", size=10)
        pdf.multi_cell(180, 10, txt=f'患者主诉: {record["patient_complaint"]}', align='L')
        pdf.multi_cell(180, 10, txt=f'病史: {record["medical_history"]}', align='L')
        pdf.multi_cell(180, 10, txt=f'体格检查: {record["physical_examination"]}', align='L')
        pdf.multi_cell(180, 10, txt=f'辅助检查: {record["auxiliary_examination"]}', align='L')
        pdf.multi_cell(180, 10, txt=f'诊断: {record["diagnosis"]}', align='L')
        pdf.multi_cell(180, 10, txt=f'治疗建议: {record["treatment_advice"]}', align='L')

        # 添加签名信息
        pdf.set_font("SimSun", '', 12)
        pdf.cell(200, 10, txt='签名信息', ln=True)
        pdf.set_font("SimSun", size=10)
        # 使用 multi_cell 方法处理长字符串自动换行
        pdf.multi_cell(180, 10, txt=f'医生签名: {record["doctor_signature"]}', align='L')
        pdf.multi_cell(180, 10, txt=f'创建时间: {record["created_at"]}', align='L')
        pdf.multi_cell(180, 10, txt=f'医生公钥: {record["doctor_public_key"]}', align='L')
        pdf.multi_cell(180, 10, txt=f'验证字符串: {record["to_cal"]}', align='L')
        pdf.multi_cell(180, 10, txt=f'验证用哈希: {record["to_verify_signature"]}', align='L')

        # 添加分页符
        pdf.add_page()

    # 保存 PDF 文档
    pdf.output(pdf_file)
    print(f"PDF 文档已保存为: {pdf_file}")


# 请给出适配这种文件格式的数字签名。如果可行，将给出的签名和哈希值直接用于数字签名，不必额外计算，验签采用sm3和sm2算法
def parse_json(file_path):
    """解析JSON文件并验证签名"""
    result = {}
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            records = json.load(f)
        # print("json_file:", records)
        for record in records:
            #print(record)
            medical_record_id = record.get('medical_record_id')
            # public_key = record.get('doctor_public_key', '')
            public_key = record.get('doctor_public_key')
            # print("json_bkey:", type(public_key), public_key)
            # print(public_key)
            # signature = record.get('to_verify_signature', '')
            signature = record.get('doctor_signature')
            # print("json_sign:", type(signature), signature)
            # print(signature)
            # message = record.get('to_cal', '')
            message = record.get('to_cal')
            # print("json_test:", type(message), message)
            # print(message)

            # signature = hexstr_bytes(signature)
            # public_key = hexstr_bytes(public_key)
            message = generate_sm3_hash(message)
            # print(message)
            if all([public_key, signature, message]):
                medical_record_id = str(medical_record_id)
                result[medical_record_id] = sm2_verify(signature, message, public_key,)
            else:
                print("验签关键字段缺失")
                result[medical_record_id] = False

        return result
    except Exception as e:
        print(f"Error parsing JSON: {e}")
        return result


def parse_docx(file_path):
    """解析DOCX文件并验证签名"""
    result = {}
    try:
        doc = docx.Document(file_path)
        record_id = None
        public_key = None
        signature = None
        message = None
        for para in doc.paragraphs:
            # print("para:", para.text)
            if '记录ID:' in para.text:
                record_id = para.text.split(':')[-1].strip()
                print("docx_mrid:", type(record_id), record_id)
            elif '验证字符串:' in para.text:
                message = para.text.split(':')[-1].strip()
                print("docx_test:", type(message), message)
            elif '医生签名:' in para.text:
                signature = para.text.split(':')[-1].strip()
                print("docx_sign:", type(signature), signature)
            elif '医生公钥:' in para.text:
                public_key = para.text.split(':')[-1].strip()
                print("docx_bkey:", type(public_key), public_key)

            if all([record_id, public_key, signature, message]):
                message = generate_sm3_hash(message)
                result[record_id] = sm2_verify(signature, message, public_key)
                # result[record_id] = False
                record_id = None
                public_key = None
                signature = None
                message = None
            else:
                # print("关键字段读取缺失")
                pass

        return result
    except Exception as e:
        print(f"Error parsing DOCX: {e}")
        return result


def parse_pdf(file_path):
    """解析PDF文件并验证签名"""
    result = {}
    record_data = {
        'record_id': None,
        'public_key': None,
        'signature': None,
        'message': None
    }
    current_field = None
    accumulated_data = ''

    try:
        with fitz.open(file_path) as doc:
            for page in doc:
                text = page.get_text()
                lines = text.split('\n')
                for line in lines:
                    line = line.strip()
                    if line.startswith('病历ID:'):
                        '''
                        # 如果已经收集到上一条记录的所有必要字段，验证并保存结果
                        if all(record_data.values()):
                            record_id = record_data['record_id']
                            public_key = record_data['public_key']
                            signature = record_data['signature']
                            message = record_data['message']
                            # 验证签名
                            result[record_id] = sm2_verify(public_key, signature, message)
                            # 重置记录数据
                            record_data = {
                                'record_id': None,
                                'public_key': None,
                                'signature': None,
                                'message': None
                            }
                            accumulated_data = ''
                        '''
                        # 收集当前行的病历ID
                        record_data['record_id'] = line.split(':', 1)[1].strip()


                    elif line.startswith('验证字符串:'):
                        # 收集验证字符串
                        current_field = 'message'
                        accumulated_data = line.split(':', 1)[1].strip()
                        record_data['message'] = accumulated_data
                    elif line.startswith('医生签名:'):
                        # 收集医生签名
                        current_field = 'signature'
                        accumulated_data = line.split(':', 1)[1].strip()
                        record_data['signature'] = accumulated_data
                    elif line.startswith('医生公钥:'):
                        # 收集医生公钥
                        current_field = 'public_key'
                        accumulated_data = line.split(':', 1)[1].strip()
                        record_data['public_key'] = accumulated_data
                    elif line.startswith('创建时间'):
                        current_field = ''
                    elif line.startswith('验证用哈希'):
                        current_field = ''
                    elif current_field:
                        # 如果当前正在收集某个字段，将当前行内容追加到该字段
                        accumulated_data += '' + line.strip()
                        if current_field == 'message':
                            record_data['message'] = accumulated_data
                        elif current_field == 'signature':
                            record_data['signature'] = accumulated_data
                        elif current_field == 'public_key':
                            record_data['public_key'] = accumulated_data
                    # 如果已经收集到所有必要字段，验证签名
                    if all(record_data.values()):
                        print(record_data)
                        record_id = record_data['record_id']
                        public_key = record_data['public_key']
                        signature = record_data['signature']
                        message = record_data['message']
                        message = generate_sm3_hash(message)
                        # 验证签名
                        result[record_id] = sm2_verify(signature, message, public_key)
                        # 重置记录数据
                        record_data = {
                            'record_id': None,
                            'public_key': None,
                            'signature': None,
                            'message': None
                        }
                        accumulated_data = ''
                        current_field = None
        return result
    except Exception as e:
        print(f"解析 PDF 错误: {e}")
        return result


def verify_signature(file_path):
    """验证文件签名的主函数"""
    if file_path.endswith('.json'):
        return parse_json(file_path)
    elif file_path.endswith('.docx'):
        return parse_docx(file_path)
    elif file_path.endswith('.pdf'):
        return parse_pdf(file_path)
    else:
        print("Unsupported file format")
        return False


if __name__ == "__main__":
    json_re = verify_signature(r"F:\2025-1spring\medical_record_file\medical_records_proposal_1_20250514_204206.json")
    print()
    docx_re = verify_signature(r"F:\2025-1spring\medical_record_file\medical_records_proposal_1_20250514_204216.docx")
    print()
    pdf_re = verify_signature(r"F:\2025-1spring\medical_record_file\medical_records_proposal_1_20250514_204225.pdf")
    print()

    print("json_re:", json_re)
    print("docx_re:", docx_re)
    print("pdf_re:", pdf_re)

    '''
    # 示例 JSON 文件路径
    json_file = r"F:\2025-1spring\250314\medical_records_proposal_3_20250509_091337.json"

    # 检查 JSON 文件是否存在
    if not os.path.exists(json_file):
        print(f"JSON 文件 '{json_file}' 不存在。")
        exit(1)

    # 读取 JSON 文件
    with open(json_file, 'r', encoding='utf-8') as f:
        json_data = json.load(f)

    # 转换为 Word 文档
    word_file = "medical_records.docx"
    json_to_word(json_data, word_file)

    # 转换为 PDF 文档
    pdf_file = "medical_records.pdf"
    json_to_pdf(json_data, pdf_file)
'''